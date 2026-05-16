"""Generate the declaration DOCX template (run once, commit output)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "wf_docx" / "templates" / "declaration.docx"
)


def main() -> None:
    _TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Source Serif Pro"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    cap = doc.add_paragraph("{{CAPTION}}")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = doc.add_paragraph("{{DECLARATION_TITLE}}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph("{{DRAFT_DATE}}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")
    doc.add_paragraph("{{SIGNATURE_BLOCK}}")

    doc.save(str(_TEMPLATE_PATH))
    print(f"Wrote {_TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
