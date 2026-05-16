"""OOXML helpers for Word comments, shading, and footnotes (python-docx gaps)."""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from datetime import UTC, datetime
from io import BytesIO
from typing import Any, Final, cast
from xml.etree import ElementTree as ET

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

_W_NS: Final[str] = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS: Final[str] = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS: Final[str] = "http://schemas.openxmlformats.org/package/2006/content-types"


@dataclass
class CommentSpec:
    comment_id: int
    start: int
    end: int
    author: str
    text: str
    initials: str = "WF"
    created_at: datetime = field(default_factory=lambda: datetime.now(tz=UTC))


@dataclass
class ShadingSpec:
    start: int
    end: int
    fill: str = "F2F2F2"


@dataclass
class FootnoteSpec:
    footnote_id: int
    start: int
    end: int
    text: str


def _clamp_span(text: str, start: int, end: int) -> tuple[int, int] | None:
    n = len(text)
    if n == 0:
        return (0, 0) if start == 0 and end == 0 else None
    s = max(0, min(start, n))
    e = max(s, min(end, n))
    return s, e


def _kind_order(kind: str) -> int:
    order = {
        "comment_start": 0,
        "shade_on": 1,
        "footnote_ref": 2,
        "shade_off": 3,
        "comment_end": 4,
    }
    return order.get(kind, 5)


def _merge_events(
    text: str,
    comments: list[CommentSpec],
    shadings: list[ShadingSpec],
    footnotes: list[FootnoteSpec],
) -> list[tuple[int, str, Any]]:
    events: list[tuple[int, str, Any]] = []
    for c in comments:
        span = _clamp_span(text, c.start, c.end)
        if span is None:
            continue
        s, e = span
        events.append((s, "comment_start", c.comment_id))
        events.append((e, "comment_end", c.comment_id))
    for sh in shadings:
        span = _clamp_span(text, sh.start, sh.end)
        if span is None:
            continue
        s, e = span
        events.append((s, "shade_on", sh.fill))
        events.append((e, "shade_off", None))
    for fn in footnotes:
        span = _clamp_span(text, fn.start, fn.end)
        if span is None:
            continue
        s, _e = span
        events.append((s, "footnote_ref", fn.footnote_id))
    events.sort(key=lambda x: (x[0], _kind_order(x[1])))
    return events


def rebuild_paragraph_with_annotations(
    paragraph: Paragraph,
    text: str,
    *,
    comments: list[CommentSpec],
    shadings: list[ShadingSpec],
    footnotes: list[FootnoteSpec],
) -> None:
    """Replace paragraph content with runs, comment ranges, shading, footnote refs."""
    p_el = paragraph._p
    for child in list(p_el):
        if child.tag in {
            qn("w:r"),
            qn("w:commentRangeStart"),
            qn("w:commentRangeEnd"),
        }:
            p_el.remove(child)

    events = _merge_events(text, comments, shadings, footnotes)
    if not events:
        if text:
            paragraph.add_run(text)
        return

    shade_depth = 0
    current_fill = "F2F2F2"
    pos = 0
    event_idx = 0

    def append_text_fragment(start: int, end: int) -> None:
        nonlocal shade_depth, current_fill
        if start >= end:
            return
        frag = text[start:end]
        if not frag:
            return
        if shade_depth > 0:
            r = paragraph.add_run(frag)
            r_pr = r._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), current_fill)
            shd.set(qn("w:color"), "auto")
            r_pr.append(shd)
        else:
            paragraph.add_run(frag)

    while event_idx < len(events) or pos < len(text):
        next_offset = len(text)
        if event_idx < len(events):
            next_offset = events[event_idx][0]
        if pos < next_offset:
            append_text_fragment(pos, next_offset)
            pos = next_offset
        if event_idx >= len(events):
            break
        offset, kind, payload = events[event_idx]
        event_idx += 1
        if offset > pos:
            append_text_fragment(pos, offset)
            pos = offset
        if kind == "comment_start":
            el = OxmlElement("w:commentRangeStart")
            el.set(qn("w:id"), str(payload))
            p_el.append(el)
        elif kind == "comment_end":
            el = OxmlElement("w:commentRangeEnd")
            el.set(qn("w:id"), str(payload))
            p_el.append(el)
            ref_run = OxmlElement("w:r")
            ref = OxmlElement("w:commentReference")
            ref.set(qn("w:id"), str(payload))
            ref_run.append(ref)
            p_el.append(ref_run)
        elif kind == "shade_on":
            shade_depth += 1
            current_fill = str(payload)
        elif kind == "shade_off":
            shade_depth = max(0, shade_depth - 1)
        elif kind == "footnote_ref":
            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            vert = OxmlElement("w:vertAlign")
            vert.set(qn("w:val"), "superscript")
            r_pr.append(vert)
            r.append(r_pr)
            fn_ref = OxmlElement("w:footnoteReference")
            fn_ref.set(qn("w:id"), str(payload))
            r.append(fn_ref)
            p_el.append(r)

    if pos < len(text):
        append_text_fragment(pos, len(text))


@dataclass
class CommentStore:
    specs: list[CommentSpec] = field(default_factory=list)
    _next_id: int = 0

    def add(
        self,
        *,
        start: int,
        end: int,
        text: str,
        author: str = "Wellfounded",
    ) -> int:
        cid = self._next_id
        self._next_id += 1
        self.specs.append(
            CommentSpec(
                comment_id=cid,
                start=start,
                end=end,
                author=author,
                text=text,
            ),
        )
        return cid


@dataclass
class FootnoteStore:
    specs: list[FootnoteSpec] = field(default_factory=list)
    bodies: dict[int, str] = field(default_factory=dict)
    _next_id: int = 1

    def add(self, *, start: int, end: int, text: str) -> int:
        fid = self._next_id
        self._next_id += 1
        self.specs.append(
            FootnoteSpec(footnote_id=fid, start=start, end=end, text=text),
        )
        self.bodies[fid] = text
        return fid


def _iso_comment_date(dt: datetime) -> str:
    return dt.astimezone(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")


def inject_comments_and_footnotes(
    docx_bytes: bytes,
    *,
    comment_store: CommentStore,
    footnote_store: FootnoteStore,
) -> bytes:
    """Add word/comments.xml and word/footnotes.xml to a saved DOCX package."""
    if not comment_store.specs and not footnote_store.bodies:
        return docx_bytes

    with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zin:
        file_map: dict[str, bytes] = {name: zin.read(name) for name in zin.namelist()}

    if comment_store.specs:
        file_map["word/comments.xml"] = _build_comments_xml(comment_store.specs)
    if footnote_store.bodies:
        file_map["word/footnotes.xml"] = _build_footnotes_xml(footnote_store.bodies)

    if comment_store.specs or footnote_store.bodies:
        file_map["[Content_Types].xml"] = _patch_content_types(
            file_map["[Content_Types].xml"],
            comments=bool(comment_store.specs),
            footnotes=bool(footnote_store.bodies),
        )
        file_map["word/_rels/document.xml.rels"] = _patch_document_rels(
            file_map["word/_rels/document.xml.rels"],
            comments=bool(comment_store.specs),
            footnotes=bool(footnote_store.bodies),
        )

    buf_out = BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in file_map.items():
            zout.writestr(name, data)
    return buf_out.getvalue()


def _build_comments_xml(specs: list[CommentSpec]) -> bytes:
    root = ET.Element(f"{{{_W_NS}}}comments")
    for spec in specs:
        c = ET.SubElement(
            root,
            f"{{{_W_NS}}}comment",
            {
                f"{{{_W_NS}}}id": str(spec.comment_id),
                f"{{{_W_NS}}}author": spec.author,
                f"{{{_W_NS}}}date": _iso_comment_date(spec.created_at),
                f"{{{_W_NS}}}initials": spec.initials,
            },
        )
        p = ET.SubElement(c, f"{{{_W_NS}}}p")
        r = ET.SubElement(p, f"{{{_W_NS}}}r")
        t = ET.SubElement(r, f"{{{_W_NS}}}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = spec.text
    return cast(bytes, ET.tostring(root, encoding="utf-8", xml_declaration=True))


def _build_footnotes_xml(bodies: dict[int, str]) -> bytes:
    root = ET.Element(f"{{{_W_NS}}}footnotes")
    for fid in sorted(bodies):
        fn = ET.SubElement(root, f"{{{_W_NS}}}footnote", {f"{{{_W_NS}}}id": str(fid)})
        p = ET.SubElement(fn, f"{{{_W_NS}}}p")
        r = ET.SubElement(p, f"{{{_W_NS}}}r")
        t = ET.SubElement(r, f"{{{_W_NS}}}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = bodies[fid]
    return cast(bytes, ET.tostring(root, encoding="utf-8", xml_declaration=True))


def _patch_content_types(
    raw: bytes,
    *,
    comments: bool,
    footnotes: bool,
) -> bytes:
    root = ET.fromstring(raw)
    overrides = {
        el.get("PartName"): el for el in root.findall(f"{{{_CT_NS}}}Override")
    }
    if comments and "/word/comments.xml" not in overrides:
        ET.SubElement(
            root,
            f"{{{_CT_NS}}}Override",
            {
                "PartName": "/word/comments.xml",
                "ContentType": (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.comments+xml"
                ),
            },
        )
    if footnotes and "/word/footnotes.xml" not in overrides:
        ET.SubElement(
            root,
            f"{{{_CT_NS}}}Override",
            {
                "PartName": "/word/footnotes.xml",
                "ContentType": (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.footnotes+xml"
                ),
            },
        )
    return cast(bytes, ET.tostring(root, encoding="utf-8", xml_declaration=True))


def _patch_document_rels(
    raw: bytes,
    *,
    comments: bool,
    footnotes: bool,
) -> bytes:
    root = ET.fromstring(raw)
    rels = list(root.findall(f"{{{_REL_NS}}}Relationship"))
    existing_targets = {el.get("Target") for el in rels}
    max_id = 0
    for el in rels:
        rid = el.get("Id", "")
        m = re.match(r"rId(\d+)", rid)
        if m:
            max_id = max(max_id, int(m.group(1)))

    def add_rel(target: str, rel_type: str) -> None:
        nonlocal max_id
        if target in existing_targets:
            return
        max_id += 1
        ET.SubElement(
            root,
            f"{{{_REL_NS}}}Relationship",
            {
                "Id": f"rId{max_id}",
                "Type": rel_type,
                "Target": target,
            },
        )

    if comments:
        add_rel(
            "comments.xml",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        )
    if footnotes:
        add_rel(
            "footnotes.xml",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )
    return cast(bytes, ET.tostring(root, encoding="utf-8", xml_declaration=True))


def count_comments_in_docx(docx_bytes: bytes) -> int:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        if "word/comments.xml" not in zf.namelist():
            return 0
        root = ET.fromstring(zf.read("word/comments.xml"))
    suffix = "}comment"
    return sum(1 for el in root.iter() if el.tag.endswith(suffix))


def document_xml_has_review_markup(docx_bytes: bytes) -> bool:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return (
        "commentRangeStart" in xml
        or "w:highlight" in xml
        or 'w:fill="F2F2F2"' in xml
        or "footnoteReference" in xml
    )
