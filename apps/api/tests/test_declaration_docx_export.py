"""Declaration draft DOCX export (unit + focused API checks)."""

from __future__ import annotations

import zipfile
from datetime import UTC, datetime
from io import BytesIO

import pytest
from docx import Document

from declarations.schemas import (
    DECLARATION_SECTION_IDS,
    DeclarationDraftContent,
    DeclarationFlag,
    DeclarationFlagStatus,
    DeclarationFlagType,
    DeclarationParagraph,
    DeclarationSection,
    FlagSpan,
    InferenceSpanOut,
    TranscriptSegment,
)
from wf_docx.declaration import (
    CleanExportBlockedError,
    DeclarationRenderInput,
    render_clean_copy,
    render_working_copy,
)
from wf_docx.ooxml_annotations import (
    count_comments_in_docx,
    document_xml_has_review_markup,
)


def _minimal_draft() -> DeclarationDraftContent:
    sections: dict[str, DeclarationSection] = {}
    for sid in DECLARATION_SECTION_IDS:
        sections[sid] = DeclarationSection(
            section_id=sid,
            paragraphs=[
                DeclarationParagraph(
                    id=f"{sid}:p0",
                    text=f"Body text for section {sid}.",
                    source_segment_ids=["seg-0"],
                ),
            ],
        )
    sections["past_persecution"] = DeclarationSection(
        section_id="past_persecution",
        paragraphs=[
            DeclarationParagraph(
                id="past_persecution:p0",
                text="In March 2023, four men detained me at my office.",
                source_segment_ids=["seg-1"],
                inference_spans=[
                    InferenceSpanOut(
                        start=40,
                        end=55,
                        rationale="Inferred location detail not stated by client.",
                    ),
                ],
            ),
        ],
    )
    return DeclarationDraftContent(sections=sections)


def _render_input(
    *,
    draft: DeclarationDraftContent,
    flags: list[DeclarationFlag],
    parallel: bool = False,
) -> DeclarationRenderInput:
    segments = [
        TranscriptSegment(
            start=0.0,
            end=10.0,
            speaker="client",
            source_text="Source segment zero.",
            english_text="English segment zero.",
        ),
        TranscriptSegment(
            start=10.0,
            end=20.0,
            speaker="client",
            source_text="Source segment one.",
            english_text="English segment one.",
        ),
    ]
    return DeclarationRenderInput(
        draft=draft,
        flags=flags,
        case_pseudonym="M.A.",
        country_code="ER",
        draft_version=1,
        generated_at=datetime(2025, 3, 1, tzinfo=UTC),
        transcript_segments=segments if parallel else None,
        parallel=parallel,
    )


def test_working_copy_has_comments_and_shading() -> None:
    draft = _minimal_draft()
    flags = [
        DeclarationFlag(
            type=DeclarationFlagType.INFERENCE,
            paragraph_id="past_persecution:p0",
            span=FlagSpan(start=0, end=20),
            description="Opening sentence inferred.",
            suggested_resolution="Confirm with client.",
        ),
        DeclarationFlag(
            type=DeclarationFlagType.GAP,
            paragraph_id="gap:first_incident_date",
            span=FlagSpan(start=0, end=0),
            description="First incident date missing.",
            suggested_resolution="Ask client for date.",
            element_key="first_incident_date",
        ),
        DeclarationFlag(
            type=DeclarationFlagType.AMBIGUITY,
            paragraph_id="past_persecution:p0",
            span=FlagSpan(start=25, end=35),
            description="Unclear who detained client.",
            suggested_resolution="Clarify identity of agents.",
        ),
    ]
    raw = render_working_copy(_render_input(draft=draft, flags=flags))
    assert len(raw) > 2000
    assert count_comments_in_docx(raw) >= 2
    assert document_xml_has_review_markup(raw)

    with zipfile.ZipFile(BytesIO(raw)) as zf:
        assert "word/comments.xml" in zf.namelist()
        doc_xml = zf.read("word/document.xml").decode("utf-8")
        assert "commentRangeStart" in doc_xml
        assert 'w:fill="F2F2F2"' in doc_xml or "F2F2F2" in doc_xml
        if any(f.type == DeclarationFlagType.AMBIGUITY for f in flags):
            assert "footnoteReference" in doc_xml
            assert "word/footnotes.xml" in zf.namelist()


def test_clean_copy_blocked_with_open_gap() -> None:
    draft = _minimal_draft()
    flags = [
        DeclarationFlag(
            type=DeclarationFlagType.GAP,
            paragraph_id="gap:first_incident_date",
            span=FlagSpan(start=0, end=0),
            description="Missing date.",
            suggested_resolution="Obtain date.",
            element_key="first_incident_date",
        ),
    ]
    with pytest.raises(CleanExportBlockedError) as exc_info:
        render_clean_copy(_render_input(draft=draft, flags=flags))
    assert exc_info.value.unresolved_flag_ids


def test_clean_copy_no_review_markup_when_resolved() -> None:
    draft = _minimal_draft()
    flags = [
        DeclarationFlag(
            type=DeclarationFlagType.GAP,
            paragraph_id="gap:first_incident_date",
            span=FlagSpan(start=0, end=0),
            description="Missing date.",
            suggested_resolution="Obtain date.",
            status=DeclarationFlagStatus.resolved,
            element_key="first_incident_date",
        ),
        DeclarationFlag(
            type=DeclarationFlagType.INFERENCE,
            paragraph_id="past_persecution:p0",
            span=FlagSpan(start=0, end=5),
            description="Inferred.",
            suggested_resolution="Confirm.",
            status=DeclarationFlagStatus.deferred,
        ),
    ]
    raw = render_clean_copy(_render_input(draft=draft, flags=flags))
    assert count_comments_in_docx(raw) == 0
    assert not document_xml_has_review_markup(raw)

    doc = Document(BytesIO(raw))
    body = " ".join(p.text for p in doc.paragraphs)
    assert "March 2023" in body


def test_parallel_layout_includes_source_and_english() -> None:
    draft = _minimal_draft()
    raw = render_working_copy(
        _render_input(draft=draft, flags=[], parallel=True),
    )
    doc = Document(BytesIO(raw))
    table_text = " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "Source segment one" in table_text
    assert "four men detained" in table_text
