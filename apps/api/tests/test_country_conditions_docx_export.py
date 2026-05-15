"""Country conditions memo DOCX export (build + API)."""

from __future__ import annotations

import re
import uuid
import zipfile
from datetime import UTC, date, datetime
from io import BytesIO

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from audit.models import AuditLogEntry
from cases.models import (
    Case,
    CaseArtifact,
    CaseArtifactType,
    CaseAssignment,
    CaseAssignmentRole,
    ClaimBasis,
)
from country_conditions.docx_memo import (
    build_country_conditions_docx_bytes,
    cited_passage_ids_ordered,
)
from country_conditions.models import CountryConditionsMemo, CountryConditionsMemoStatus
from country_conditions.schemas import (
    CC_SECTION_IDS,
    CountryConditionsInputs,
    FinalMemoStructured,
    MemoSectionStructured,
)
from orgs.models import UserRole, UserStatus
from orgs.repository import OrgRepository
from retrieval.schemas import PassageExportMeta


def _cite(pid: uuid.UUID) -> str:
    return f'<cite passage_id="{pid}"/>'


def _final_memo_two_cites(pid_a: uuid.UUID, pid_b: uuid.UUID) -> FinalMemoStructured:
    sections: list[MemoSectionStructured] = []
    for i, sid in enumerate(CC_SECTION_IDS):
        if i == 0:
            body = f"Intro {_cite(pid_a)} and {_cite(pid_b)}."
        elif i == 1:
            body = f"Second section repeats {_cite(pid_a)}."
        else:
            body = f"Section {sid} narrative without cites."
        sections.append(
            MemoSectionStructured(
                section_id=sid,
                title=f"Heading for {sid}",
                body=body,
            ),
        )
    return FinalMemoStructured(
        sections=sections,
        bibliography=[
            {
                "index": 1,
                "passage_id": str(pid_a),
                "source_title": "Doc A",
                "publication_date": "2020-01-01",
                "url": "https://example.invalid/a",
                "section_anchor": "x",
            },
            {
                "index": 2,
                "passage_id": str(pid_b),
                "source_title": "Doc B",
                "publication_date": "2021-01-01",
                "url": "https://example.invalid/b",
                "section_anchor": "y",
            },
        ],
    )


def _export_meta(
    pid_a: uuid.UUID,
    pid_b: uuid.UUID,
) -> dict[uuid.UUID, PassageExportMeta]:
    lv = datetime(2024, 6, 1, 12, 0, 0, tzinfo=UTC)
    return {
        pid_a: PassageExportMeta(
            passage_id=pid_a,
            source_family="state_dept_human_rights",
            document_title="State report A",
            publication_date=date(2020, 1, 15),
            url="https://example.invalid/a",
            section_anchor="s1",
            last_verified_at=lv,
        ),
        pid_b: PassageExportMeta(
            passage_id=pid_b,
            source_family="amnesty",
            document_title="Amnesty report B",
            publication_date=date(2021, 3, 1),
            url="https://example.invalid/b",
            section_anchor="s2",
            last_verified_at=lv,
        ),
    }


def test_build_docx_structure_and_citation_indices() -> None:
    pid_a = uuid.uuid4()
    pid_b = uuid.uuid4()
    final = _final_memo_two_cites(pid_a, pid_b)
    meta = _export_meta(pid_a, pid_b)
    inputs = CountryConditionsInputs(
        country_code="ER",
        basis=ClaimBasis.political_opinion,
        group_description="Test group",
        timeframe_start_year=2020,
        jurisdiction_asylum_office=None,
    )
    raw = build_country_conditions_docx_bytes(
        final=final,
        inputs=inputs,
        case_pseudonym="M.A.",
        export_meta=meta,
        memo_generated_at=datetime(2025, 1, 2, tzinfo=UTC),
    )
    assert len(raw) > 2000

    doc = Document(BytesIO(raw))
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2")
    assert h2_count == 6

    ordered = cited_passage_ids_ordered(final)
    assert ordered == [pid_a, pid_b]

    with zipfile.ZipFile(BytesIO(raw)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "<cite" not in xml
    assert xml.count("w:hyperlink") >= 4
    supers = len(re.findall(r'w:val="superscript"', xml))
    assert supers >= 3

    bib_para = next(
        (p for p in doc.paragraphs if p.text.startswith("1. ") and "Amnesty" in p.text),
        None,
    )
    assert bib_para is not None
    assert "Amnesty report B" in bib_para.text
    p2 = next(p for p in doc.paragraphs if p.text.startswith("2. "))
    assert "State report A" in p2.text


@pytest.mark.asyncio(loop_scope="session")
async def test_export_docx_endpoint_audit(
    db_session: AsyncSession,
    api_client: AsyncClient,
    capsys: pytest.CaptureFixture[str],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pid_a = uuid.UUID("aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa")
    pid_b = uuid.UUID("bbbbbbbb-bbbb-4bbb-8bbb-bbbbbbbbbbbb")
    final = _final_memo_two_cites(pid_a, pid_b)
    meta = _export_meta(pid_a, pid_b)

    async def _fake_fetch(
        _session: AsyncSession,
        passage_ids: list[uuid.UUID],
    ) -> dict[uuid.UUID, PassageExportMeta]:
        return {pid: meta[pid] for pid in passage_ids}

    monkeypatch.setattr("cases.router.fetch_passages_export_meta", _fake_fetch)

    org_repo = OrgRepository(db_session)
    slug = f"dx-{uuid.uuid4().hex[:8]}"
    org = await org_repo.create_org(name="DX Org", slug=slug)
    u = await org_repo.create_user(
        organization_id=org.id,
        email="dx@example.com",
        display_name="DX",
        role=UserRole.attorney,
        status=UserStatus.active,
    )
    await db_session.commit()

    await api_client.post(
        "/auth/magic-link",
        json={"email": "dx@example.com", "organization_slug": slug},
    )
    out = capsys.readouterr().out
    m = re.search(r"token=([^\s]+)", out)
    assert m is not None
    await api_client.get(f"/auth/callback?token={m.group(1)}", follow_redirects=False)

    case = Case(
        id=uuid.uuid4(),
        organization_id=org.id,
        pseudonym="M.A.",
        country_code="ER",
        basis=ClaimBasis.political_opinion,
        group_description="G",
        filing_deadline=None,
        asylum_office=None,
        intake_notes="n",
        created_by_user_id=u.id,
    )
    db_session.add(case)
    db_session.add(
        CaseAssignment(
            case_id=case.id,
            user_id=u.id,
            role_on_case=CaseAssignmentRole.lead_attorney,
        ),
    )
    await db_session.flush()
    art_id = uuid.uuid4()
    memo_id = uuid.uuid4()
    db_session.add(
        CaseArtifact(
            id=art_id,
            case_id=case.id,
            artifact_type=CaseArtifactType.country_conditions_memo,
        ),
    )
    db_session.add(
        CountryConditionsMemo(
            id=memo_id,
            organization_id=org.id,
            case_id=case.id,
            case_artifact_id=art_id,
            status=CountryConditionsMemoStatus.complete,
            inputs=CountryConditionsInputs(
                country_code="ER",
                basis=ClaimBasis.political_opinion,
                group_description="G",
                timeframe_start_year=2020,
                jurisdiction_asylum_office=None,
            ).model_dump(mode="json"),
            output=final.model_dump(mode="json"),
            version=2,
            generated_by_user_id=u.id,
            generated_at=datetime(2025, 1, 1, tzinfo=UTC),
            model_versions={},
            error_message=None,
            correlation_request_id=uuid.uuid4(),
        ),
    )
    await db_session.commit()

    r = await api_client.get(
        f"/cases/{case.id}/country-conditions/{memo_id}/export.docx",
    )
    assert r.status_code == 200, r.text
    assert (
        r.headers.get("content-type", "").split(";")[0].strip()
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(r.content) > 1000
    assert "attachment" in r.headers.get("content-disposition", "")

    res = await db_session.execute(
        select(AuditLogEntry).where(
            AuditLogEntry.action == "country_conditions.memo.export.docx",
            AuditLogEntry.resource_id == memo_id,
        ),
    )
    assert res.scalar_one_or_none() is not None
